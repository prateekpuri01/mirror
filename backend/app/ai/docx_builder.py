"""Build tailored .docx resumes from LLM-generated content.

Style resolution precedence (highest first):
  1. ``resume_design`` selection passed to ``build_docx``
  2. ``user_profile.data["resume_design"]``
  3. ``docs/resume_style.yaml`` (power-user override)
  4. ``_DEFAULT_STYLE`` below

Layout, color scheme, and font come from preset partials in
``resume_presets.py``. Style is resolved per-request and threaded through a
``BuildContext`` so swapping designs at runtime is safe.
"""

import logging
import os
import re
from dataclasses import dataclass, field
from datetime import UTC, datetime
from pathlib import Path
from types import SimpleNamespace
from typing import Any

import yaml
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Inches, Pt, RGBColor, Twips

from app.ai.resume_presets import resolve_style

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Default style (used when no preset / YAML / param overrides apply)
# ---------------------------------------------------------------------------

_DEFAULT_STYLE: dict[str, Any] = {
    "colors": {
        "navy": "#1A1A1A",  # near-black — name (bold), section headers
        "orange": "#B8851E",  # mustard — the pop: tagline, accents, rules
        "contact": "#555555",  # mid-gray — contact line
        "separator": "#B8B8B8",  # light gray — separator characters
        "dark": "#1F1F1F",  # near-black — body text
        "name_first": "#8B7355",  # warm tan — first name
        "link": "#8B6914",  # darker mustard — readable links on white
        "border": "#B8851E",  # mustard — section-header underline POP
    },
    "fonts": {
        "name_light": "Cambria",
        "name_bold": "Cambria",
        "body": "Aptos",
        "skill_label": "Aptos",
    },
    "sizes": {
        "name": 30,
        "tagline": 10.5,
        "contact": 9,
        "section_header": 12,
        "body": 10,
        "body_small": 9.5,
    },
    "margins": {
        "top": 0.6,
        "bottom": 0.6,
        "left": 0.7,
        "right": 0.7,
    },
    "layout": {
        "kind": "banner",
        "columns": 1,
        "header_kind": "banner",
        "section_style": "spaced",
        "section_underline": False,
    },
}


OUTPUT_DIR = "/app/output/resumes"

# Line spacing constants (independent of font choice)
LINE_SPACING_BODY = Emu(248 * 12700 // 240)
LINE_SPACING_COMPACT = Emu(244 * 12700 // 240)


# ---------------------------------------------------------------------------
# Style loading & context
# ---------------------------------------------------------------------------


def _load_yaml_style() -> dict:
    """Load ``docs/resume_style.yaml`` over ``_DEFAULT_STYLE`` if present.

    Returns the default style untouched if no YAML file exists or it fails
    to parse.
    """
    for path in (Path("/app/docs/resume_style.yaml"), Path("docs/resume_style.yaml")):
        if not path.exists():
            continue
        try:
            with open(path) as f:
                user_style = yaml.safe_load(f) or {}
            merged: dict[str, Any] = {}
            for section, default_block in _DEFAULT_STYLE.items():
                if isinstance(default_block, dict):
                    merged[section] = {**default_block, **(user_style.get(section) or {})}
                else:
                    merged[section] = user_style.get(section, default_block)
            logger.info("Loaded resume style from %s", path)
            return merged
        except Exception:
            logger.warning("Failed to load resume style from %s, using defaults", path)
    return {k: (dict(v) if isinstance(v, dict) else v) for k, v in _DEFAULT_STYLE.items()}


def _load_style_for_request(
    explicit_design: dict | None,
    profile_data: dict | None,
) -> dict:
    """Resolve the style dict for one ``build_docx`` call.

    Precedence (highest first):
      1. ``explicit_design`` (param passed to ``build_docx``)
      2. ``profile_data["resume_design"]``
      3. YAML overrides (already folded into the base style)
      4. ``_DEFAULT_STYLE``

    If no design selection is present at any level, returns the YAML/default
    base unchanged — preserves backward compat with the maintainer's
    ``resume_style.yaml`` (e.g. Mustard theme).
    """
    base_style = _load_yaml_style()

    selection = explicit_design
    if not selection and profile_data:
        selection = profile_data.get("resume_design")

    if not selection:
        # Make sure the base style has a layout block.
        if "layout" not in base_style:
            base_style["layout"] = dict(_DEFAULT_STYLE["layout"])
        return base_style

    return resolve_style(selection, base_style)


def _hex_to_rgb(hex_str: str) -> RGBColor:
    h = hex_str.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def _inches_to_emu(inches: float) -> int:
    return int(inches * 914400)


@dataclass
class BuildContext:
    """Resolved style + computed values, threaded through every render helper.

    Holds the raw style dict alongside pre-computed RGBColor/Pt/Emu values
    so render code can use readable accessors (``ctx.color.navy``) instead
    of re-parsing hex strings at every call site.
    """

    style: dict
    color: SimpleNamespace = field(default_factory=SimpleNamespace)
    font: SimpleNamespace = field(default_factory=SimpleNamespace)
    size: SimpleNamespace = field(default_factory=SimpleNamespace)
    margin: SimpleNamespace = field(default_factory=SimpleNamespace)
    border_hex: str = ""
    layout: dict = field(default_factory=dict)


def _make_context(style: dict) -> BuildContext:
    colors = style["colors"]
    fonts = style["fonts"]
    sizes = style["sizes"]
    margins = style["margins"]
    layout = style.get("layout") or dict(_DEFAULT_STYLE["layout"])

    return BuildContext(
        style=style,
        color=SimpleNamespace(
            navy=_hex_to_rgb(colors["navy"]),
            orange=_hex_to_rgb(colors["orange"]),
            contact=_hex_to_rgb(colors["contact"]),
            separator=_hex_to_rgb(colors["separator"]),
            dark=_hex_to_rgb(colors["dark"]),
            name_first=_hex_to_rgb(colors["name_first"]),
            link=_hex_to_rgb(colors["link"]),
        ),
        font=SimpleNamespace(
            name_light=fonts["name_light"],
            name_bold=fonts["name_bold"],
            body=fonts["body"],
            skill_label=fonts["skill_label"],
        ),
        size=SimpleNamespace(
            name=Pt(sizes["name"]),
            tagline=Pt(sizes["tagline"]),
            contact=Pt(sizes["contact"]),
            section_header=Pt(sizes["section_header"]),
            body=Pt(sizes["body"]),
            body_small=Pt(sizes["body_small"]),
        ),
        margin=SimpleNamespace(
            top=Emu(_inches_to_emu(margins["top"])),
            bottom=Emu(_inches_to_emu(margins["bottom"])),
            left=Emu(_inches_to_emu(margins["left"])),
            right=Emu(_inches_to_emu(margins["right"])),
        ),
        border_hex=colors["border"].lstrip("#"),
        layout=layout,
    )


# ---------------------------------------------------------------------------
# OOXML helpers (style-independent)
# ---------------------------------------------------------------------------


def _add_bottom_border(paragraph, sz: str = "6", space: str = "2", color: str = "B8851E") -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), sz)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_paragraph_top_border(
    paragraph, sz: str = "6", space: str = "1", color: str = "B8851E"
) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), sz)
    top.set(qn("w:space"), space)
    top.set(qn("w:color"), color)
    pBdr.append(top)
    pPr.append(pBdr)


def _set_table_width(table, width_in: float) -> None:
    """Pin a table's total width to a specific inch value (in OOXML twips)."""
    tblPr = table._element.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        table._element.insert(0, tblPr)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(int(width_in * 1440)))
    tblW.set(qn("w:type"), "dxa")


def _set_cell_left_border(cell, color: str, sz: str = "12") -> None:
    """Set ONLY a left border on a cell, clearing the others.

    Used for the timeline layout's vertical accent rule running down the
    experience block.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove any existing tcBorders we might have set earlier
    for existing in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(existing)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), sz)
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), color.lstrip("#"))
    tcBorders.append(left)
    tcPr.append(tcBorders)


def _set_section_columns(section, num: int, space_twips: int = 360) -> None:
    """Configure a section's column count via native Word ``w:cols``.

    Native column flow is what ATS parsers expect; tables and text boxes
    are the layouts that fail to parse, so we avoid both.
    """
    sectPr = section._sectPr
    cols = sectPr.find(qn("w:cols"))
    if cols is None:
        cols = OxmlElement("w:cols")
        sectPr.append(cols)
    cols.set(qn("w:num"), str(num))
    cols.set(qn("w:space"), str(space_twips))
    cols.set(qn("w:equalWidth"), "1")


def _insert_column_break(doc) -> None:
    """Force the boundary between columns deterministically."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "column")
    run._r.append(br)


# ---------------------------------------------------------------------------
# Run-level + paragraph helpers
# ---------------------------------------------------------------------------


def _set_run(
    ctx: BuildContext,
    run,
    *,
    font_name: str | None = None,
    size=None,
    color=None,
    bold: bool = False,
    italic: bool = False,
    underline: bool = False,
) -> None:
    run.font.name = font_name or ctx.font.body
    run.font.size = size or ctx.size.body
    if color is not None:
        run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    if underline:
        run.font.underline = True


def _add_hyperlink(ctx: BuildContext, paragraph, url: str, text: str) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), ctx.font.body)
    rFonts.set(qn("w:hAnsi"), ctx.font.body)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "19")  # 9.5pt in half-points
    rPr.append(sz)

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_elem.append(t)

    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _add_bullet_paragraph(ctx: BuildContext, doc, text: str, font_size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Emu(19050)
    p.paragraph_format.space_after = Emu(31750)
    p.paragraph_format.left_indent = Twips(360)
    p.paragraph_format.first_line_indent = Twips(-200)

    bullet_run = p.add_run("•  ")
    _set_run(ctx, bullet_run, size=font_size or ctx.size.body_small)

    run = p.add_run(text)
    _set_run(ctx, run, size=font_size or ctx.size.body_small)
    return p


# ---------------------------------------------------------------------------
# Section renderers (container-agnostic — they append to whatever section
# the doc is currently writing into, which the layout dispatcher controls).
# ---------------------------------------------------------------------------


def _build_contact_items(personal: dict) -> list[tuple[str, str, str | None]]:
    """Build the (emoji, display_text, url-or-None) tuples for the contact line.

    Shared between the centered, banner, and split (compact) headers.
    """
    items: list[tuple[str, str, str | None]] = []
    phone = personal.get("phone", "")
    email = personal.get("email", "")
    linkedin = personal.get("linkedin", "")
    if phone:
        digits = "".join(c for c in phone if c.isdigit())
        formatted = f"{digits[:3]}-{digits[3:6]}-{digits[6:]}" if len(digits) == 10 else phone
        items.append(("\U0001f4f1 ", formatted, None))
    if email:
        items.append(("✉ ", email, f"mailto:{email}"))
    if linkedin:
        linkedin_url = linkedin if linkedin.startswith("http") else f"https://{linkedin}"
        display = (
            linkedin_url.replace("https://www.", "")
            .replace("http://www.", "")
            .replace("https://", "")
            .replace("http://", "")
            .rstrip("/")
        )
        items.append(("\U0001f310 ", display, linkedin_url))
    return items


def _build_professional_links(personal: dict) -> list[tuple[str, str]]:
    scholar = personal.get("google_scholar", "")
    professional_links: list[dict] = personal.get("professional_links", [])
    links: list[tuple[str, str]] = []
    if scholar:
        scholar_url = scholar if scholar.startswith("http") else f"https://{scholar}"
        links.append(("Google Scholar", scholar_url))
    for pl in professional_links:
        url = pl.get("url", "")
        label = pl.get("label", "")
        if url and label:
            if not url.startswith("http"):
                url = f"https://{url}"
            links.append((label, url))
    return links


def _render_contact_line(
    ctx: BuildContext,
    paragraph,
    items: list[tuple[str, str, str | None]],
) -> None:
    """Render emoji/text/url contact items into an existing paragraph."""
    for i, (emoji, text, url) in enumerate(items):
        if i > 0:
            run = paragraph.add_run("    ")
            _set_run(
                ctx, run, font_name=ctx.font.body, size=ctx.size.contact, color=ctx.color.contact
            )
        run = paragraph.add_run(emoji)
        _set_run(ctx, run, font_name=ctx.font.body, size=ctx.size.contact, color=ctx.color.contact)
        if url:
            _add_hyperlink(ctx, paragraph, url, text)
        else:
            run = paragraph.add_run(text)
            _set_run(
                ctx, run, font_name=ctx.font.body, size=ctx.size.contact, color=ctx.color.contact
            )


def _build_contact_grid_items(personal: dict) -> list[dict]:
    """Items for the timeline designer header's contact grid (1-4 cells).

    Each: ``{"icon": str, "value": str, "url": str | None}``. Order:
    location, phone, email, linkedin (capped at 4 to keep the grid balanced).
    """
    items: list[dict] = []
    location = personal.get("location", "")
    phone = personal.get("phone", "")
    email = personal.get("email", "")
    linkedin = personal.get("linkedin", "")
    if location:
        items.append({"icon": "📍", "value": location, "url": None})
    if phone:
        digits = "".join(c for c in phone if c.isdigit())
        formatted = f"{digits[:3]}-{digits[3:6]}-{digits[6:]}" if len(digits) == 10 else phone
        items.append({"icon": "📞", "value": formatted, "url": None})
    if email:
        items.append({"icon": "✉", "value": email, "url": f"mailto:{email}"})
    if linkedin:
        linkedin_url = linkedin if linkedin.startswith("http") else f"https://{linkedin}"
        display = (
            linkedin_url.replace("https://www.", "")
            .replace("http://www.", "")
            .replace("https://", "")
            .replace("http://", "")
            .rstrip("/")
        )
        items.append({"icon": "🔗", "value": display, "url": linkedin_url})
    return items[:4]


def _add_hyperlink_white(
    ctx: BuildContext, paragraph, url: str, text: str, size_pt: float = 9.0
) -> None:
    """Hyperlink styled for the dark band header — light gray, no underline."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run_elem = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), ctx.font.body)
    rFonts.set(qn("w:hAnsi"), ctx.font.body)
    rPr.append(rFonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size_pt * 2)))
    rPr.append(sz)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "E6E6E6")
    rPr.append(color)
    run_elem.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_elem.append(t)
    hyperlink.append(run_elem)
    paragraph._p.append(hyperlink)


def _render_timeline_designer_header(
    ctx: BuildContext, doc, profile_data: dict, tagline: str
) -> None:
    """Designer header for the timeline layout — colored band, 3 stacked rows.

    Row 1: centered name in large white reverse type with letter-spacing.
    Row 2: centered tagline flanked by hairline rules.
    Row 3: contact grid (1-4 cells, icon over value, vertical white dividers).
    """
    personal = profile_data.get("personal", {})
    full_name = personal.get("name", "")
    grid_items = _build_contact_grid_items(personal)
    n_cols = max(1, len(grid_items))

    table_width_in = 8.5 - (ctx.style["margins"]["left"] + ctx.style["margins"]["right"])
    accent_hex = ctx.style["colors"]["navy"].lstrip("#")

    table = doc.add_table(rows=3, cols=n_cols)
    table.autofit = False
    _set_table_width(table, table_width_in)

    # Shade every cell in the band and clear borders.
    col_in = table_width_in / n_cols
    for row in table.rows:
        for cell in row.cells:
            _set_cell_shading(cell, accent_hex)
            _remove_cell_borders(cell)
            cell.width = Inches(col_in)

    # --- Row 1: name (merged) ---
    name_cell = table.cell(0, 0)
    if n_cols > 1:
        name_cell = name_cell.merge(table.cell(0, n_cols - 1))
    _set_cell_margins(name_cell, top_in=0.26, right_in=0.30, bottom_in=0.06, left_in=0.30)
    name_p = name_cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(0)
    run = name_p.add_run(full_name.upper())
    _set_run(
        ctx,
        run,
        font_name=ctx.font.name_bold,
        size=Pt(26),
        color=RGBColor(0xFF, 0xFF, 0xFF),
        bold=True,
    )
    _set_run_char_spacing(run, 60)  # 3pt letter expansion

    # --- Row 2: tagline with hairline flanking rules (merged) ---
    tag_cell = table.cell(1, 0)
    if n_cols > 1:
        tag_cell = tag_cell.merge(table.cell(1, n_cols - 1))
    _set_cell_margins(tag_cell, top_in=0.04, right_in=0.30, bottom_in=0.18, left_in=0.30)
    tag_p = tag_cell.paragraphs[0]
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tag_p.paragraph_format.space_before = Pt(0)
    tag_p.paragraph_format.space_after = Pt(0)
    if tagline:
        # Tighter sizing so long taglines (~60+ chars) fit on one line within
        # the printable width. Tagline at 8.5pt with 1.5pt letter spacing,
        # rules shortened to 5 chars each.
        rule = "─" * 5
        run_l = tag_p.add_run(rule + "  ")
        _set_run(
            ctx,
            run_l,
            font_name=ctx.font.body,
            size=Pt(8),
            color=RGBColor(0xCC, 0xCC, 0xCC),
        )
        run_t = tag_p.add_run(tagline.upper())
        _set_run(
            ctx,
            run_t,
            font_name=ctx.font.body,
            size=Pt(8.5),
            color=RGBColor(0xFF, 0xFF, 0xFF),
            bold=True,
        )
        _set_run_char_spacing(run_t, 30)
        run_r = tag_p.add_run("  " + rule)
        _set_run(
            ctx,
            run_r,
            font_name=ctx.font.body,
            size=Pt(8),
            color=RGBColor(0xCC, 0xCC, 0xCC),
        )

    # --- Row 3: contact grid ---
    for j, item in enumerate(grid_items):
        cell = table.cell(2, j)
        cell.width = Inches(col_in)
        _set_cell_margins(cell, top_in=0.10, right_in=0.10, bottom_in=0.18, left_in=0.10)
        if j > 0:
            # Thin white vertical separator between grid cells.
            _set_cell_left_border(cell, color="FFFFFF", sz="4")

        icon_p = cell.paragraphs[0]
        icon_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        icon_p.paragraph_format.space_before = Pt(0)
        icon_p.paragraph_format.space_after = Pt(1)
        run = icon_p.add_run(item["icon"])
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=Pt(13),
            color=RGBColor(0xFF, 0xFF, 0xFF),
        )

        val_p = cell.add_paragraph()
        val_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        val_p.paragraph_format.space_before = Pt(0)
        val_p.paragraph_format.space_after = Pt(0)
        if item.get("url"):
            _add_hyperlink_white(ctx, val_p, item["url"], item["value"], size_pt=8.0)
        else:
            run = val_p.add_run(item["value"])
            _set_run(
                ctx,
                run,
                font_name=ctx.font.body,
                size=Pt(8),
                color=RGBColor(0xE6, 0xE6, 0xE6),
            )

    # Professional links below the band — keep readable as normal hyperlinks.
    links = _build_professional_links(personal)
    if links:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(2)
        for i, (label, url) in enumerate(links):
            if i > 0:
                run = p.add_run("  |  ")
                _set_run(
                    ctx,
                    run,
                    font_name=ctx.font.body,
                    size=ctx.size.contact,
                    color=ctx.color.separator,
                )
            _add_hyperlink(ctx, p, url, label)


def _render_centered_header(ctx: BuildContext, doc, profile_data: dict, tagline: str) -> None:
    """Centered name + tagline rule + contact line. Used by the timeline layout."""
    personal = profile_data.get("personal", {})
    full_name = personal.get("name", "")
    name_parts = full_name.split(" ", 1)
    first_name = name_parts[0] if name_parts else ""
    last_name = name_parts[1] if len(name_parts) > 1 else ""

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)

    if first_name:
        run = p.add_run(first_name.upper())
        _set_run(
            ctx, run, font_name=ctx.font.name_light, size=ctx.size.name, color=ctx.color.name_first
        )
    if last_name:
        spacer = p.add_run(" ")
        spacer.font.size = ctx.size.name
        run = p.add_run(last_name.upper())
        _set_run(
            ctx,
            run,
            font_name=ctx.font.name_bold,
            size=ctx.size.name,
            color=ctx.color.dark,
            bold=True,
        )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(tagline)
    _set_run(ctx, run, font_name=ctx.font.body, size=ctx.size.tagline, color=ctx.color.orange)
    _add_bottom_border(p, sz="2", space="4", color=ctx.border_hex)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(1)
    _render_contact_line(ctx, p, _build_contact_items(personal))

    links = _build_professional_links(personal)
    if links:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for i, (label, url) in enumerate(links):
            if i > 0:
                run = p.add_run("  |  ")
                _set_run(
                    ctx,
                    run,
                    font_name=ctx.font.body,
                    size=ctx.size.contact,
                    color=ctx.color.separator,
                )
            _add_hyperlink(ctx, p, url, label)


def _render_banner_header(ctx: BuildContext, doc, profile_data: dict, tagline: str) -> None:
    """Filled-color band header: name in white reverse type, tagline below.

    The banner is a 1-cell shaded table spanning the printable width. The
    contact line and professional links render below the band as normal
    centered paragraphs.
    """
    personal = profile_data.get("personal", {})
    full_name = personal.get("name", "")

    table_width_in = 8.5 - (ctx.style["margins"]["left"] + ctx.style["margins"]["right"])
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    _set_table_width(table, table_width_in)

    cell = table.cell(0, 0)
    cell.width = Inches(table_width_in)
    accent_hex = ctx.style["colors"]["navy"].lstrip("#")
    _set_cell_shading(cell, accent_hex)
    _remove_cell_borders(cell)
    _set_cell_margins(cell, top_in=0.22, right_in=0.32, bottom_in=0.22, left_in=0.32)

    name_p = cell.paragraphs[0]
    name_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(0)
    run = name_p.add_run(full_name.upper())
    _set_run(
        ctx,
        run,
        font_name=ctx.font.name_bold,
        size=Pt(24),
        color=RGBColor(0xFF, 0xFF, 0xFF),
        bold=True,
    )

    if tagline:
        tag_p = cell.add_paragraph()
        tag_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        tag_p.paragraph_format.space_before = Pt(2)
        tag_p.paragraph_format.space_after = Pt(0)
        run = tag_p.add_run(tagline)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=Pt(11),
            color=RGBColor(0xE6, 0xE6, 0xE6),
        )

    # Contact line below the band — centered, normal style.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    _render_contact_line(ctx, p, _build_contact_items(personal))

    links = _build_professional_links(personal)
    if links:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for i, (label, url) in enumerate(links):
            if i > 0:
                run = p.add_run("  |  ")
                _set_run(
                    ctx,
                    run,
                    font_name=ctx.font.body,
                    size=ctx.size.contact,
                    color=ctx.color.separator,
                )
            _add_hyperlink(ctx, p, url, label)


def _render_compact_header(ctx: BuildContext, doc, profile_data: dict, tagline: str) -> None:
    """Split header: name + tagline left, contact stacked & right-aligned.

    A borderless 2-column table holds the layout. Followed by a thin
    accent rule across the page.
    """
    personal = profile_data.get("personal", {})
    full_name = personal.get("name", "")

    table_width_in = 8.5 - (ctx.style["margins"]["left"] + ctx.style["margins"]["right"])
    name_col_in = table_width_in * 0.6
    contact_col_in = table_width_in - name_col_in

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    _set_table_width(table, table_width_in)

    name_cell = table.cell(0, 0)
    contact_cell = table.cell(0, 1)
    name_cell.width = Inches(name_col_in)
    contact_cell.width = Inches(contact_col_in)
    _remove_cell_borders(name_cell)
    _remove_cell_borders(contact_cell)
    _set_cell_margins(name_cell, top_in=0.0, right_in=0.1, bottom_in=0.05, left_in=0.0)
    _set_cell_margins(contact_cell, top_in=0.04, right_in=0.0, bottom_in=0.05, left_in=0.1)

    # Name (left)
    name_p = name_cell.paragraphs[0]
    name_p.paragraph_format.space_before = Pt(0)
    name_p.paragraph_format.space_after = Pt(0)
    run = name_p.add_run(full_name)
    _set_run(
        ctx,
        run,
        font_name=ctx.font.name_bold,
        size=Pt(20),
        color=ctx.color.navy,
        bold=True,
    )

    if tagline:
        tag_p = name_cell.add_paragraph()
        tag_p.paragraph_format.space_before = Pt(1)
        tag_p.paragraph_format.space_after = Pt(0)
        run = tag_p.add_run(tagline)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=Pt(10.5),
            color=ctx.color.orange,
            italic=True,
        )

    # Contact (right, right-aligned, stacked one per line)
    contact_items = _build_contact_items(personal)
    for i, (emoji, text, url) in enumerate(contact_items):
        p = contact_cell.paragraphs[0] if i == 0 else contact_cell.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(emoji)
        _set_run(ctx, run, font_name=ctx.font.body, size=ctx.size.contact, color=ctx.color.contact)
        if url:
            _add_hyperlink(ctx, p, url, text)
        else:
            run = p.add_run(text)
            _set_run(
                ctx, run, font_name=ctx.font.body, size=ctx.size.contact, color=ctx.color.contact
            )

    # Thin accent rule below header
    rule_p = doc.add_paragraph()
    rule_p.paragraph_format.space_before = Pt(4)
    rule_p.paragraph_format.space_after = Pt(2)
    _add_paragraph_top_border(rule_p, sz="6", space="0", color=ctx.border_hex)

    links = _build_professional_links(personal)
    if links:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        for i, (label, url) in enumerate(links):
            if i > 0:
                run = p.add_run("  |  ")
                _set_run(
                    ctx,
                    run,
                    font_name=ctx.font.body,
                    size=ctx.size.contact,
                    color=ctx.color.separator,
                )
            _add_hyperlink(ctx, p, url, label)


def _render_section_header(ctx: BuildContext, doc, title: str) -> None:
    """Dispatch to the right section-header style based on layout config.

    Styles: ``rule`` (default — bold caps with bottom border), ``spaced``
    (letter-spaced caps in accent color, no rule), ``accent_bar`` (colored
    block prefix + bold caps, no rule).
    """
    style = ctx.layout.get("section_style", "rule")
    if style == "spaced":
        return _render_section_header_spaced(ctx, doc, title)
    if style == "accent_bar":
        return _render_section_header_accent_bar(ctx, doc, title)
    # default = "rule"
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Emu(127000)
    p.paragraph_format.space_after = Emu(38100)
    run = p.add_run(title.upper())
    _set_run(
        ctx,
        run,
        font_name=ctx.font.body,
        size=ctx.size.section_header,
        color=ctx.color.navy,
        bold=True,
    )
    if ctx.layout.get("section_underline", True):
        _add_bottom_border(p, sz="6", space="2", color=ctx.border_hex)


def _set_run_char_spacing(run, twentieths_of_point: int) -> None:
    """Add character spacing (kerning) via OOXML ``w:spacing``.

    Twentieths of a point — so ``40`` = 2pt expansion between every glyph.
    Word keeps the natural word-space width on top of this, so word gaps
    stay visibly wider than letter gaps (the bug with the old
    string-interleaving approach).
    """
    rPr = run._r.get_or_add_rPr()
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:val"), str(twentieths_of_point))
    rPr.append(spacing)


def _render_section_header_spaced(ctx: BuildContext, doc, title: str) -> None:
    """Letter-spaced caps for the banner layout — real OOXML kerning."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    _set_run(
        ctx,
        run,
        font_name=ctx.font.body,
        size=ctx.size.section_header,
        color=ctx.color.orange,
        bold=True,
    )
    _set_run_char_spacing(run, 60)


def _render_section_header_accent_bar(ctx: BuildContext, doc, title: str) -> None:
    """Colored block prefix + bold caps. Used by the compact layout."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    bar = p.add_run("▍ ")
    _set_run(
        ctx,
        bar,
        font_name=ctx.font.body,
        size=Pt(13),
        color=ctx.color.orange,
        bold=True,
    )
    run = p.add_run(title.upper())
    _set_run(
        ctx,
        run,
        font_name=ctx.font.body,
        size=ctx.size.section_header,
        color=ctx.color.navy,
        bold=True,
    )


def _render_summary(ctx: BuildContext, doc, summary: str) -> None:
    _render_section_header(ctx, doc, "Summary")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(summary)
    _set_run(ctx, run, size=ctx.size.body)


def _render_selected_research(
    ctx: BuildContext,
    doc,
    research_entries: list,
    section_title: str = "Selected Research",
) -> None:
    _render_section_header(ctx, doc, section_title)
    for entry in research_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(0)

        label = entry.get("category_label", "RESEARCH")
        run = p.add_run(f"{label.upper()} — ")
        _set_run(
            ctx,
            run,
            font_name=ctx.font.skill_label,
            size=ctx.size.body,
            color=ctx.color.orange,
            bold=True,
        )

        run = p.add_run(entry.get("title", ""))
        _set_run(
            ctx, run, font_name=ctx.font.body, size=ctx.size.body, color=ctx.color.navy, bold=True
        )

        desc = entry.get("description", "")
        if desc:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.25)
            run = p.add_run(desc)
            _set_run(ctx, run, size=ctx.size.body_small)


def _render_experience_block(
    ctx: BuildContext,
    doc,
    org: str,
    title: str,
    location: str,
    dates: str,
    bullets: list,
) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(org)
    _set_run(ctx, run, font_name=ctx.font.body, size=ctx.size.body, color=ctx.color.navy, bold=True)

    if location:
        run = p.add_run(f"  —  {location}")
        _set_run(ctx, run, size=ctx.size.body, color=ctx.color.contact)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title)
    _set_run(ctx, run, size=ctx.size.body, italic=True)

    if dates:
        run = p.add_run(f"  |  {dates}")
        _set_run(ctx, run, size=ctx.size.body_small, color=ctx.color.separator)

    for bullet in bullets:
        bullet_text = bullet["text"] if isinstance(bullet, dict) else bullet
        _add_bullet_paragraph(ctx, doc, bullet_text)


def _render_experience(ctx: BuildContext, doc, experience_data: dict, profile_data: dict) -> None:
    from app.ai.utils import employer_key

    _render_section_header(ctx, doc, "Professional Experience")
    work_history = profile_data.get("work_history", [])

    wh_by_key = {}
    for wh in work_history:
        key = employer_key(wh.get("employer", ""))
        end = wh.get("end") or "Present"
        wh_by_key[key] = {
            "org": wh.get("employer", ""),
            "title": wh.get("title", ""),
            "location": wh.get("location", ""),
            "dates": f"{wh.get('start', '')} – {end}",
        }

    for wh in work_history:
        key = employer_key(wh.get("employer", ""))
        emp_data = experience_data.get(key, {})
        bullets = emp_data.get("bullets", [])
        if not bullets:
            continue
        info = wh_by_key.get(key, {})
        _render_experience_block(
            ctx,
            doc,
            info.get("org", key),
            info.get("title", ""),
            info.get("location", ""),
            info.get("dates", ""),
            bullets,
        )

    rendered_keys = {employer_key(wh.get("employer", "")) for wh in work_history}
    for key, emp_data in experience_data.items():
        if key in rendered_keys:
            continue
        bullets = emp_data.get("bullets", [])
        if bullets:
            _render_experience_block(ctx, doc, key, "", "", "", bullets)


def _render_publications(ctx: BuildContext, doc, publications: list) -> None:
    if not publications:
        return
    _render_section_header(ctx, doc, "Selected Publications")
    for pub in publications:
        citation = pub.get("citation", "")
        if not citation:
            continue
        _add_bullet_paragraph(ctx, doc, citation, font_size=ctx.size.body_small)


def _render_skills(ctx: BuildContext, doc, skills_data: dict) -> None:
    _render_section_header(ctx, doc, "Technical Skills")
    category_labels = {
        "ai_systems": "AI Systems",
        "data_science": "Data Science",
        "engineering": "Engineering",
    }
    for key, label in category_labels.items():
        value = skills_data.get(key, "")
        if not value:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Emu(19050)
        p.paragraph_format.space_after = Emu(31750)
        run = p.add_run(f"{label}: ")
        _set_run(
            ctx,
            run,
            font_name=ctx.font.skill_label,
            size=ctx.size.body_small,
            color=ctx.color.navy,
            bold=True,
        )
        run = p.add_run(value)
        _set_run(ctx, run, size=ctx.size.body_small)


def _render_education(ctx: BuildContext, doc, profile_data: dict) -> None:
    education = profile_data.get("education", [])
    if not education:
        return
    _render_section_header(ctx, doc, "Education")
    for edu in education:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Emu(6350)
        p.paragraph_format.space_after = Emu(6350)
        p.paragraph_format.left_indent = Emu(63500)

        degree_field = f"{edu.get('degree', '')} {edu.get('field', '')}".strip()
        institution = edu.get("institution", "")
        year = edu.get("year", "")
        honors = edu.get("honors", "")

        run = p.add_run(f"{degree_field}, {institution}, {year}")
        _set_run(ctx, run, size=ctx.size.body)
        if honors:
            run = p.add_run(f" — {honors}")
            _set_run(ctx, run, size=ctx.size.body, italic=True)


def _render_awards(ctx: BuildContext, doc, awards_text: str) -> None:
    if not awards_text:
        return
    _render_section_header(ctx, doc, "Awards & Honors")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(awards_text)
    _set_run(ctx, run, size=ctx.size.body_small)


# ---------------------------------------------------------------------------
# Layout dispatch
# ---------------------------------------------------------------------------

LAYOUT_DEFAULT_ORDER: dict[str, list[str]] = {
    "banner":         ["summary", "selected_research", "experience", "publications",
                       "technical_skills", "education", "awards"],
    "compact":        ["summary", "selected_research", "experience", "publications",
                       "technical_skills", "education", "awards"],
    "centered_clean": ["summary", "selected_research", "experience", "publications",
                       "technical_skills", "education", "awards"],
    "timeline":       ["summary", "experience_education", "selected_research",
                       "technical_skills", "publications", "awards"],
    "two_column":     ["summary", "experience", "selected_research", "publications", "awards"],
}


def _render_ordered_sections(ctx, container, resume_data, profile_data, layout):
    order = resume_data.get("section_order") or LAYOUT_DEFAULT_ORDER.get(layout, LAYOUT_DEFAULT_ORDER["banner"])
    for section_id in order:
        if section_id == "summary":
            _render_summary(ctx, container, resume_data.get("summary", ""))
        elif section_id == "experience":
            _render_experience(ctx, container, resume_data.get("experience", {}), profile_data)
        elif section_id == "experience_education":
            _render_experience_education_timeline(ctx, container, resume_data.get("experience", {}), profile_data)
        elif section_id == "selected_research":
            _render_selected_research(
                ctx,
                container,
                resume_data.get("selected_research", []),
                section_title=resume_data.get(
                    "selected_research_section_title", "Selected Research"
                ),
            )
        elif section_id == "publications":
            _render_publications(ctx, container, resume_data.get("publications", []))
        elif section_id == "technical_skills":
            _render_skills(ctx, container, resume_data.get("technical_skills", {}))
        elif section_id == "education":
            _render_education(ctx, container, profile_data)
        elif section_id == "awards":
            _render_awards(ctx, container, resume_data.get("awards", ""))

def _build_layout_banner(ctx: BuildContext, doc, resume_data: dict, profile_data: dict) -> None:
    _render_banner_header(ctx, doc, profile_data, resume_data.get("tagline", ""))
    _render_ordered_sections(ctx, doc, resume_data, profile_data, "banner")


def _build_layout_compact(ctx: BuildContext, doc, resume_data: dict, profile_data: dict) -> None:
    _render_compact_header(ctx, doc, profile_data, resume_data.get("tagline", ""))
    _render_ordered_sections(ctx, doc, resume_data, profile_data, "compact")


def _build_layout_centered_clean(
    ctx: BuildContext, doc, resume_data: dict, profile_data: dict
) -> None:
    """Clean centered layout — no colored band, no date gutter.

    Pairs the two-tone centered header (``_render_centered_header``: light-gray
    first name + dark last name + orange tagline with hairline underline) with
    the banner's body structure (separate Experience and Education sections).
    Good fit for academic / policy CVs where the designer band feels too
    heavy and the timeline's combined experience+education table doesn't
    match the user's intent.
    """
    _render_centered_header(ctx, doc, profile_data, resume_data.get("tagline", ""))
    _render_ordered_sections(ctx, doc, resume_data, profile_data, "centered_clean")


def _build_layout_timeline(ctx: BuildContext, doc, resume_data: dict, profile_data: dict) -> None:
    _render_timeline_designer_header(ctx, doc, profile_data, resume_data.get("tagline", ""))
    _render_ordered_sections(ctx, doc, resume_data, profile_data, "timeline")


def _render_experience_education_timeline(
    ctx: BuildContext,
    doc,
    experience_data: dict,
    profile_data: dict,
) -> None:
    """Combined work + education timeline.

    Title-first ordering: bold job/degree title on the first line, italic
    employer-or-institution-and-location on the second, bullets below. A
    single 2-col table renders both work and education rows so the vertical
    accent rule is continuous; an italic "Education" sub-row separates the
    two sub-groups without breaking the rule.
    """
    from app.ai.utils import employer_key

    _render_section_header(ctx, doc, "Experience & Education")
    work_history = profile_data.get("work_history", [])
    education = profile_data.get("education", [])

    # ---- Work blocks -----------------------------------------------------
    work_blocks: list[dict] = []
    for wh in work_history:
        key = employer_key(wh.get("employer", ""))
        emp_data = experience_data.get(key, {})
        bullets = emp_data.get("bullets", [])
        if not bullets:
            continue
        start = wh.get("start", "")
        end = wh.get("end") or "Present"
        # Normalize a couple of trailing-year noise patterns
        date_str = f"{start} — {end}" if start else end
        employer = wh.get("employer", key)
        location = wh.get("location", "")
        subtitle = f"{employer}  ·  {location}" if location else employer
        work_blocks.append(
            {
                "kind": "work",
                "date": date_str,
                "title": wh.get("title", ""),
                "subtitle": subtitle,
                "bullets": bullets,
            }
        )

    rendered = {employer_key(wh.get("employer", "")) for wh in work_history}
    for key, emp in experience_data.items():
        if key in rendered:
            continue
        bullets = emp.get("bullets", [])
        if not bullets:
            continue
        work_blocks.append(
            {
                "kind": "work",
                "date": "",
                "title": key,
                "subtitle": "",
                "bullets": bullets,
            }
        )

    # ---- Education blocks ------------------------------------------------
    edu_blocks: list[dict] = []
    for edu in education:
        year = str(edu.get("year", ""))
        degree_field = f"{edu.get('degree', '')} {edu.get('field', '')}".strip()
        institution = edu.get("institution", "")
        honors = edu.get("honors", "")
        if institution and honors:
            subtitle = f"{institution}  —  {honors}"
        else:
            subtitle = institution or honors
        edu_blocks.append(
            {
                "kind": "education",
                "date": year,
                "title": degree_field,
                "subtitle": subtitle,
                "bullets": [],
            }
        )

    # ---- Assemble blocks, inserting Education sub-heading row ------------
    blocks: list[dict] = list(work_blocks)
    if edu_blocks:
        blocks.append({"kind": "subheading", "label": "Education"})
        blocks.extend(edu_blocks)

    if not blocks:
        return

    gutter_in = ctx.layout.get("date_gutter_in", 0.95)
    total_in = 8.5 - (ctx.style["margins"]["left"] + ctx.style["margins"]["right"])
    content_in = total_in - gutter_in

    table = doc.add_table(rows=len(blocks), cols=2)
    table.autofit = False
    _set_table_width(table, total_in)

    for i, block in enumerate(blocks):
        date_cell = table.cell(i, 0)
        content_cell = table.cell(i, 1)
        date_cell.width = Inches(gutter_in)
        content_cell.width = Inches(content_in)
        _remove_cell_borders(date_cell)
        # Continuous left border on content cells = uninterrupted vertical rule.
        _set_cell_left_border(content_cell, color=ctx.border_hex, sz="12")

        if block["kind"] == "subheading":
            # Sub-heading row: small italic label, tighter padding, no date.
            _set_cell_margins(date_cell, top_in=0.0, right_in=0.18, bottom_in=0.0, left_in=0.0)
            _set_cell_margins(content_cell, top_in=0.10, right_in=0.0, bottom_in=0.04, left_in=0.22)
            sub_p = content_cell.paragraphs[0]
            sub_p.paragraph_format.space_before = Pt(0)
            sub_p.paragraph_format.space_after = Pt(0)
            run = sub_p.add_run(block["label"])
            _set_run(
                ctx,
                run,
                font_name=ctx.font.body,
                size=ctx.size.body_small,
                color=ctx.color.contact,
                italic=True,
                bold=True,
            )
            continue

        _set_cell_margins(date_cell, top_in=0.08, right_in=0.18, bottom_in=0.18, left_in=0.0)
        _set_cell_margins(content_cell, top_in=0.06, right_in=0.0, bottom_in=0.18, left_in=0.22)

        # Date — single line e.g. "2022 — Now" or "2018"
        date_p = date_cell.paragraphs[0]
        date_p.paragraph_format.space_before = Pt(0)
        date_p.paragraph_format.space_after = Pt(0)
        run = date_p.add_run(block["date"])
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=ctx.size.body_small,
            color=ctx.color.navy,
            bold=True,
        )

        # Title (bold) on first line.
        title_p = content_cell.paragraphs[0]
        title_p.paragraph_format.space_before = Pt(0)
        title_p.paragraph_format.space_after = Pt(0)
        run = title_p.add_run(block["title"])
        _set_run(
            ctx, run, font_name=ctx.font.body, size=ctx.size.body, color=ctx.color.navy, bold=True
        )

        if block["subtitle"]:
            sub_p = content_cell.add_paragraph()
            sub_p.paragraph_format.space_before = Pt(0)
            sub_p.paragraph_format.space_after = Pt(2)
            run = sub_p.add_run(block["subtitle"])
            _set_run(
                ctx,
                run,
                font_name=ctx.font.body,
                size=ctx.size.body_small,
                color=ctx.color.contact,
                italic=True,
            )

        for bullet in block["bullets"]:
            bullet_text = bullet["text"] if isinstance(bullet, dict) else bullet
            _add_bullet_paragraph(ctx, content_cell, bullet_text)


def _set_cell_shading(cell, hex_color: str) -> None:
    """Fill a table cell with a solid background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _remove_cell_borders(cell) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "nil")
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _set_cell_margins(
    cell, top_in: float, right_in: float, bottom_in: float, left_in: float
) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (
        ("top", top_in),
        ("left", left_in),
        ("bottom", bottom_in),
        ("right", right_in),
    ):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"), str(int(val * 1440)))  # inches → twips
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _render_sidebar_section_header(ctx: BuildContext, cell, title: str) -> None:
    """Section header for the dark sidebar: white text, thin white rule."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    _set_run(
        ctx,
        run,
        font_name=ctx.font.body,
        size=ctx.size.section_header,
        color=RGBColor(0xFF, 0xFF, 0xFF),
        bold=True,
    )
    _add_bottom_border(p, sz="4", space="2", color="FFFFFF")


def _render_sidebar_name(ctx: BuildContext, cell, profile_data: dict) -> None:
    personal = profile_data.get("personal", {})
    full_name = personal.get("name", "")
    name_parts = full_name.split(" ", 1)
    first_name = name_parts[0] if name_parts else ""
    last_name = name_parts[1] if len(name_parts) > 1 else ""

    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    if first_name:
        run = p.add_run(first_name)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.name_bold,
            size=Pt(22),
            color=RGBColor(0xFF, 0xFF, 0xFF),
            bold=True,
        )
    if last_name:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        run = p2.add_run(last_name)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.name_bold,
            size=Pt(22),
            color=RGBColor(0xFF, 0xFF, 0xFF),
            bold=True,
        )


def _render_sidebar_contact(ctx: BuildContext, cell, profile_data: dict) -> None:
    personal = profile_data.get("personal", {})
    email = personal.get("email", "")
    phone = personal.get("phone", "")
    linkedin = personal.get("linkedin", "")
    location = personal.get("location", "")

    _render_sidebar_section_header(ctx, cell, "Contact")
    items: list[tuple[str, str]] = []
    if location:
        items.append(("Address", location))
    if phone:
        digits = "".join(c for c in phone if c.isdigit())
        items.append(
            ("Phone", f"{digits[:3]}-{digits[3:6]}-{digits[6:]}" if len(digits) == 10 else phone)
        )
    if email:
        items.append(("E-mail", email))
    if linkedin:
        link = linkedin.replace("https://", "").replace("http://", "").rstrip("/")
        items.append(("LinkedIn", link))

    for label, value in items:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(label)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=ctx.size.body_small,
            color=RGBColor(0xFF, 0xFF, 0xFF),
            bold=True,
        )
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        run = p2.add_run(value)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=ctx.size.body_small,
            color=RGBColor(0xE6, 0xE6, 0xE6),
        )


def _render_sidebar_skills(ctx: BuildContext, cell, skills_data: dict) -> None:
    category_labels = {
        "ai_systems": "AI Systems",
        "data_science": "Data Science",
        "engineering": "Engineering",
    }
    items: list[str] = []
    for key in category_labels:
        value = skills_data.get(key, "")
        if value:
            # Sidebar bullets are flatter — split comma-separated skills.
            for s in (s.strip() for s in value.split(",")):
                if s:
                    items.append(s)
    if not items:
        return
    _render_sidebar_section_header(ctx, cell, "Skills")
    for item in items[:12]:  # keep sidebar scannable; trims very long lists
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.12)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        run = p.add_run("•  ")
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=ctx.size.body_small,
            color=RGBColor(0xFF, 0xFF, 0xFF),
        )
        run = p.add_run(item)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=ctx.size.body_small,
            color=RGBColor(0xE6, 0xE6, 0xE6),
        )


def _render_sidebar_education(ctx: BuildContext, cell, profile_data: dict) -> None:
    education = profile_data.get("education", [])
    if not education:
        return
    _render_sidebar_section_header(ctx, cell, "Education")
    for edu in education:
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(0)
        degree_field = f"{edu.get('degree', '')} {edu.get('field', '')}".strip()
        run = p.add_run(degree_field)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=ctx.size.body_small,
            color=RGBColor(0xFF, 0xFF, 0xFF),
            bold=True,
        )
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        meta = ", ".join(filter(None, [edu.get("institution", ""), str(edu.get("year", ""))]))
        run = p2.add_run(meta)
        _set_run(
            ctx,
            run,
            font_name=ctx.font.body,
            size=ctx.size.body_small,
            color=RGBColor(0xE6, 0xE6, 0xE6),
        )


def _build_layout_two_column(ctx: BuildContext, doc, resume_data: dict, profile_data: dict) -> None:
    """Modern two-column with a filled sidebar.

    Trades native-column ATS safety for the colored-sidebar look — most
    designer-template resumes use this shape. The UI surfaces this trade-off
    via an "ATS warning" chip on the design tab.

    Implementation: a single full-width table with two cells (sidebar +
    main). The sidebar cell is shaded with the accent color; main cell is
    plain white. Borders are removed so the table is invisible.
    """
    # Outer table sized to the printable width (page width minus margins).
    table_width_in = 8.5 - (ctx.style["margins"]["left"] + ctx.style["margins"]["right"])
    sidebar_width_in = ctx.layout.get("sidebar_width_in", 2.3)
    main_width_in = table_width_in - sidebar_width_in

    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.allow_autofit = False

    # Pin column widths
    sidebar_cell = table.cell(0, 0)
    main_cell = table.cell(0, 1)
    sidebar_cell.width = Inches(sidebar_width_in)
    main_cell.width = Inches(main_width_in)
    # Also set on the underlying tblGrid so Word respects them
    for cell, _w_in in ((sidebar_cell, sidebar_width_in), (main_cell, main_width_in)):
        for _tc in cell._tc.findall(qn("w:tc")):
            pass  # noqa: PIE790 — width set above is sufficient

    # Sidebar shading + margins
    sidebar_hex = ctx.style["colors"]["navy"].lstrip("#")
    _set_cell_shading(sidebar_cell, sidebar_hex)
    _set_cell_margins(sidebar_cell, top_in=0.25, right_in=0.22, bottom_in=0.25, left_in=0.25)
    _remove_cell_borders(sidebar_cell)

    # Main cell margins + borders
    _set_cell_margins(main_cell, top_in=0.20, right_in=0.20, bottom_in=0.25, left_in=0.30)
    _remove_cell_borders(main_cell)

    # --- Sidebar contents -------------------------------------------------
    _render_sidebar_name(ctx, sidebar_cell, profile_data)
    _render_sidebar_contact(ctx, sidebar_cell, profile_data)
    _render_sidebar_skills(ctx, sidebar_cell, resume_data.get("technical_skills", {}))
    _render_sidebar_education(ctx, sidebar_cell, profile_data)

    # --- Main contents ----------------------------------------------------
    # Tagline acts as the top summary line on the right
    tagline = resume_data.get("tagline", "")
    if tagline:
        p = main_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(tagline)
        _set_run(
            ctx, run, font_name=ctx.font.body, size=ctx.size.body, color=ctx.color.dark, italic=True
        )

    _render_ordered_sections(ctx, main_cell, resume_data, profile_data, "two_column")


_LAYOUT_DISPATCH = {
    "banner": _build_layout_banner,
    "compact": _build_layout_compact,
    "two_column": _build_layout_two_column,
    "timeline": _build_layout_timeline,
    "centered_clean": _build_layout_centered_clean,
}


# ---------------------------------------------------------------------------
# Top-level builder
# ---------------------------------------------------------------------------


def _sanitize_filename(s: str) -> str:
    s = re.sub(r'[<>:"/\\|?*]', "", s)
    s = re.sub(r"\s+", " ", s.strip())
    return s


def build_docx(
    resume_data: dict,
    profile_data: dict,
    job_id: str,
    company: str | None = None,
    title: str | None = None,
    resume_design: dict | None = None,
) -> str:
    """Build a formatted .docx resume.

    ``resume_design`` (when set) overrides ``profile_data["resume_design"]``,
    which itself overrides any ``docs/resume_style.yaml`` overrides.
    """
    style = _load_style_for_request(resume_design, profile_data)
    ctx = _make_context(style)

    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = ctx.margin.top
    section.bottom_margin = ctx.margin.bottom
    section.left_margin = ctx.margin.left
    section.right_margin = ctx.margin.right

    normal_style = doc.styles["Normal"]
    normal_style.font.name = ctx.font.body
    normal_style.font.size = ctx.size.body
    normal_style.paragraph_format.space_before = Pt(0)
    normal_style.paragraph_format.space_after = Pt(0)

    layout_kind = ctx.layout.get("kind", "banner")
    build_fn = _LAYOUT_DISPATCH.get(layout_kind)
    if build_fn is None:
        logger.warning("Unknown layout kind %r — falling back to banner", layout_kind)
        build_fn = _build_layout_banner
    build_fn(ctx, doc, resume_data, profile_data)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    if company and title:
        personal = profile_data.get("personal", {})
        name = personal.get("name", "Resume")
        safe_name = _sanitize_filename(name)
        safe_company = _sanitize_filename(company)
        safe_title = _sanitize_filename(title)
        filename = f"{safe_name} {safe_company} {safe_title}.docx"
    else:
        timestamp = datetime.now(UTC).strftime("%Y%m%d_%H%M%S")
        filename = f"{job_id}_{timestamp}.docx"
    filepath = os.path.join(OUTPUT_DIR, filename)

    doc.save(filepath)
    logger.info("Resume saved to %s", filepath)
    return filepath
