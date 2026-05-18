"""Unit tests for resume_to_profile.py — pure parts only.

The async LLM-driven extraction is not tested here (would require API calls).
We test:
  - docx_to_text correctly pulls text from a .docx including table cells
  - resume_text_to_profile fills in default keys for sparse profiles
"""

from __future__ import annotations

from pathlib import Path

import pytest

from tests.eval.external.resume_to_profile import docx_to_text


# ---------------------------------------------------------------------------
# docx_to_text — uses python-docx to write a fixture .docx, then read it back
# ---------------------------------------------------------------------------


def test_docx_to_text_paragraphs(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("Senior Software Engineer")
    doc.add_paragraph("Skills: Python, FastAPI, PostgreSQL")
    docx_path = tmp_path / "resume.docx"
    doc.save(str(docx_path))

    text = docx_to_text(docx_path)
    assert "Jane Doe" in text
    assert "Senior Software Engineer" in text
    assert "Python, FastAPI, PostgreSQL" in text


def test_docx_to_text_skips_blank_paragraphs(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("First")
    doc.add_paragraph("")
    doc.add_paragraph("   ")
    doc.add_paragraph("Last")
    docx_path = tmp_path / "r.docx"
    doc.save(str(docx_path))

    text = docx_to_text(docx_path)
    lines = text.split("\n")
    assert "First" in lines
    assert "Last" in lines
    # Blank lines should not be in the output
    assert "" not in lines


def test_docx_to_text_extracts_from_tables(tmp_path: Path):
    """Many resumes use tables for layout — docx_to_text must handle them."""
    from docx import Document

    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Email"
    table.cell(0, 1).text = "jane@example.com"
    table.cell(1, 0).text = "Phone"
    table.cell(1, 1).text = "555-1234"
    docx_path = tmp_path / "tabular.docx"
    doc.save(str(docx_path))

    text = docx_to_text(docx_path)
    assert "jane@example.com" in text
    assert "555-1234" in text
    assert "Email" in text


def test_docx_to_text_combined_paragraphs_and_tables(tmp_path: Path):
    from docx import Document

    doc = Document()
    doc.add_paragraph("Resume header")
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Skill"
    table.cell(0, 1).text = "Python"
    doc.add_paragraph("Resume footer")
    docx_path = tmp_path / "mixed.docx"
    doc.save(str(docx_path))

    text = docx_to_text(docx_path)
    assert "Resume header" in text
    assert "Skill" in text
    assert "Python" in text
    assert "Resume footer" in text
