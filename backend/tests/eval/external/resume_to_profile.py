"""Convert raw resume text into a UserProfile-compatible dict for scoring.

Reuses backend/app/ai/resume_extractor.py::extract_profile_from_resume to get
the structured profile, then post-processes into the dict shape that
scoring._build_compact_profile expects.
"""

from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def docx_to_text(docx_path: str | Path) -> str:
    """Extract plain text from a .docx file using python-docx."""
    from docx import Document

    doc = Document(str(docx_path))
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    # Also pull text from tables (resumes commonly use them for layout)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
    return "\n".join(parts)


async def resume_text_to_profile(resume_text: str) -> dict:
    """Run the live resume extractor and return a profile dict.

    The returned dict is compatible with scoring._build_compact_profile —
    same shape as the user's UserProfile.data field.
    """
    from app.ai.resume_extractor import extract_profile_from_resume

    profile = await extract_profile_from_resume(resume_text)
    # Ensure the keys scoring expects are present (extractor populates most of these
    # but a few may be missing for sparse resumes — fill in safe defaults)
    profile.setdefault("personal", {})
    profile.setdefault("target_roles", [])
    profile.setdefault("domains", [])
    profile.setdefault("skills", {"technical": [], "communication": [], "tools": []})
    profile.setdefault("work_history", [])
    profile.setdefault("education", [])
    profile.setdefault("awards", [])
    profile.setdefault("experience_years", "N/A")
    profile.setdefault("search_preferences", {})
    profile.setdefault("complete_profile", {"accomplishments": [], "publications": []})
    return profile


async def docx_to_profile(docx_path: str | Path) -> dict:
    """End-to-end: .docx file → profile dict ready for score_pair()."""
    text = docx_to_text(docx_path)
    if not text.strip():
        raise ValueError(f"No text extracted from {docx_path}")
    return await resume_text_to_profile(text)


async def text_blob_to_profile(text: str) -> dict:
    """For HF-style datasets where the resume is a plain text blob."""
    if not text.strip():
        raise ValueError("Empty resume text")
    return await resume_text_to_profile(text)
